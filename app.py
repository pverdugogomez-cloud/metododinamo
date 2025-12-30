import streamlit as st
import matplotlib.pyplot as plt
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re
import json
import os

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Diagnóstico Dínamo", page_icon="📊", layout="wide")

# --- GESTIÓN DEL CONTADOR ---
ARCHIVO_CONTADOR = "contador_global.json"

def obtener_y_actualizar_contador():
    # Intenta leer el archivo, si no existe, empieza en 0
    if not os.path.exists(ARCHIVO_CONTADOR):
        datos = {"total_reportes": 0}
        with open(ARCHIVO_CONTADOR, "w") as f:
            json.dump(datos, f)
    
    with open(ARCHIVO_CONTADOR, "r+") as f:
        try:
            datos = json.load(f)
        except:
            datos = {"total_reportes": 0}
            
        datos["total_reportes"] += 1
        f.seek(0)
        json.dump(datos, f)
        f.truncate()
        
    return datos["total_reportes"]

def leer_contador_actual():
    if os.path.exists(ARCHIVO_CONTADOR):
        with open(ARCHIVO_CONTADOR, "r") as f:
            try:
                return json.load(f)["total_reportes"]
            except:
                return 0
    return 0

# --- BASE DE DATOS (PREGUNTAS) ---
db_encuesta = {
    "PILAR COGNITIVO": [
        {"pregunta": "1. Estado de tu energía mental cotidiana", "opciones": ["Me siento mentalmente apagado/a o saturado/a", "Funciono, pero con esfuerzo y dispersión", "Me siento claro/a y funcional la mayor parte del día", "Me siento activo/a mentalmente, creativo/a y con ideas"]},
        {"pregunta": "2. Actividades que encienden tu mente", "opciones": ["Me cuesta y las evito", "Las hago, pero sin mayor disfrute", "Me activan y me ayudan a concentrarme", "Me energizan y sacan lo mejor de mí"]},
        {"pregunta": "3. Capacidad de enfoque", "opciones": ["Muy bajo, me distraigo con facilidad", "Intermitente, depende del día", "Bueno, logro enforcarme cuando lo necesito", "Alto, entro fácilmente en estados de concentración profunda"]},
        {"pregunta": "4. Sensación de crecimiento cognitivo", "opciones": ["Siento que estoy estancado", "Siento poco avance", "Siento que estoy creciendo gradualmente", "Siento que estoy expandiendo activamente mi potencial"]},
        {"pregunta": "5. Desarrollo de pasatiempos", "opciones": ["No tengo pasatiempos o los he dejado de lado", "Tengo intereses, pero los practico muy poco", "Tengo al menos un pasatiempo que practico con cierta regularidad", "Mis pasatiempos son una fuente clara de energía y motivación"]},
        {"pregunta": "6. Impacto de los pasatiempos en tu potencial", "opciones": ["No noto mayor efecto", "Me distraen, pero no me activan demasiado", "Me ayudan a desconectarme y sentirme mejor", "Siento que activan habilidades y aspectos valiosos de mí"]},
        {"pregunta": "7. Consumo de redes sociales", "opciones": ["Siento que me consumen mucha energía y atención", "Las uso bastante y me cuesta regularlas", "Las uso con cierto control", "Las uso de manera consciente y no interfieren con mi foco"]},
        {"pregunta": "8. Efecto de las RRSS en tu mente", "opciones": ["Me dispersan y me saturan", "A veces me afectan, a veces no", "Generalmente no interfieren", "Mantengo claridad y foco independiente de su uso"]}
    ],
    "PILAR FÍSICO": [
        {"pregunta": "9. Nivel de energía física diaria", "opciones": ["Baja, me siento cansado/a gran parte del tiempo", "Variable, con subidas y bajadas", "Adecuada para mis actividades diarias", "Alta, me siento vital y con impulso"]},
        {"pregunta": "10. Relación con tu cuerpo", "opciones": ["Desconectada, lo siento como una carga", "Funcional, pero poco consciente", "Bastante conectada y respetuosa", "Muy conectada, lo siento como un aliado"]},
        {"pregunta": "11. Movimiento y actividad corporal", "opciones": ["Está prácticamente ausente", "Aparece de forma esporádica", "Está presente de manera regular", "Es un pilar que potencia mi energía"]},
        {"pregunta": "12. Alimentación equilibrada", "opciones": ["Siento que desordena o afecta negativamente mi energía", "Es irregular y poco consciente", "Es mayormente equilibrada", "Es un hábito que potencia mi bienestar y energía"]},
        {"pregunta": "13. Impacto de la alimentación en tu funcionamiento", "opciones": ["Me genera cansancio, malestar o baja energía", "A veces afecta, a veces me sostiene", "Generalmente me sostiene bien", "Claramente mejora mi energía, concentración y ánimo"]},
        {"pregunta": "14. Higiene de sueño (Calidad)", "opciones": ["Poco reparador, despierto cansado/a", "Irregular, con noches buenas y malas", "Bastante reparador", "Profundo y reparador"]},
        {"pregunta": "15. Higiene de sueño (Hábitos)", "opciones": ["No tengo rutinas y duermo desordenadamente", "Tengo algunos intentos, pero soy irregular", "Mantengo rutinas relativamente estables", "Tengo hábitos de sueños claros y establecidos"]}
    ],
    "PILAR ESPIRITUAL": [
        {"pregunta": "16. Conexión contigo mismo/a", "opciones": ["Bajo, me siento desconectado/a de mí", "Intermitente", "Presente la mayor parte del tiempo", "Profundo y estable"]},
        {"pregunta": "17. Sentido y coherencia vital", "opciones": ["Siento poca coherencia o sentido", "Siento dudas frecuentes", "Siento bastante coherencia", "Siento que estoy alineado/a con quien soy"]},
        {"pregunta": "18. Práctica de pausas conscientes o meditación", "opciones": ["No realizo ninguna práctica", "Las realizo de forma muy esporádica", "Las realizo con cierta regularidad", "Son un pilar importante para mi bienestar"]},
        {"pregunta": "19. Efecto de estas prácticas en tu vida", "opciones": ["No están presentes, por lo tanto, no influyen", "Me ayudan solo en momentos puntuales", "Me ayudan a regularme emocionalmente", "Me ayudan a mantener coherencia, calma y claridad"]},
        {"pregunta": "20. Relaciones significativas", "opciones": ["Me drenan o me generar conflicto", "Son neutras, cumplen una función", "Me aportan apoyo y contención", "Potencian lo mejor de mí"]},
        {"pregunta": "21. Conexión con algo más grande", "opciones": ["Inexistente", "Esporádica", "Presente", "Profunda y nutritiva"]},
        {"pregunta": "22. Sensación global de activación personal", "opciones": ["Estás sobreviviendo", "Estás funcionando", "Estás en proceso de activación", "Estás desplegando tu potencial"]}
    ]
}

# --- PLAN DE ACCIÓN ---
db_plan_accion = {
    "PILAR COGNITIVO": {
        "BAJO": [
            {"accion": "Desconexión Digital: Apagar pantallas 1h antes de dormir.", "frecuencia": "Diario (Noche)"},
            {"accion": "Técnica Pomodoro: 25 min de trabajo enfocado.", "frecuencia": "2 veces al día"},
            {"accion": "Lectura ligera (no trabajo) por 15 minutos.", "frecuencia": "Fines de Semana"}
        ],
        "MEDIO": [
            {"accion": "Bloque de Enfoque Profundo (sin celular).", "frecuencia": "Diario (90 min)"},
            {"accion": "Aprender algo nuevo (podcast, video educativo).", "frecuencia": "3 veces por semana"},
            {"accion": "Limpieza de redes sociales (dejar de seguir cuentas tóxicas).", "frecuencia": "Mensual"}
        ],
        "ALTO": [
            {"accion": "Enseñar o mentorear a alguien en tu área.", "frecuencia": "Semanal"},
            {"accion": "Escribir ideas creativas o journaling.", "frecuencia": "Diario (Mañana)"},
            {"accion": "Desafío intelectual complejo (idioma, ajedrez).", "frecuencia": "Fines de Semana"}
        ]
    },
    "PILAR FÍSICO": {
        "BAJO": [
            {"accion": "Hidratación consciente (1 vaso al despertar).", "frecuencia": "Diario"},
            {"accion": "Caminata suave de 15 minutos.", "frecuencia": "Diario"},
            {"accion": "Establecer hora fija para ir a la cama.", "frecuencia": "Diario"}
        ],
        "MEDIO": [
            {"accion": "Ejercicio de fuerza o resistencia media.", "frecuencia": "3 veces por semana"},
            {"accion": "Preparar comidas saludables (Meal Prep).", "frecuencia": "Semanal (Domingo)"},
            {"accion": "Pausa activa de estiramiento.", "frecuencia": "Cada 2 horas"}
        ],
        "ALTO": [
            {"accion": "Entrenamiento de alta intensidad o deporte.", "frecuencia": "4-5 veces por semana"},
            {"accion": "Optimización del sueño (temperatura, oscuridad total).", "frecuencia": "Diario"},
            {"accion": "Actividad física en la naturaleza (trekking, etc).", "frecuencia": "Mensual"}
        ]
    },
    "PILAR ESPIRITUAL": {
        "BAJO": [
            {"accion": "Respiración consciente (3 minutos).", "frecuencia": "Diario (Al despertar)"},
            {"accion": "Contacto con la naturaleza (parque, jardín).", "frecuencia": "Semanal"},
            {"accion": "Evitar noticias negativas en la mañana.", "frecuencia": "Diario"}
        ],
        "MEDIO": [
            {"accion": "Diario de Gratitud (3 cosas buenas del día).", "frecuencia": "Diario (Noche)"},
            {"accion": "Conversación profunda con un amigo/familiar.", "frecuencia": "Semanal"},
            {"accion": "Práctica de meditación guiada (10 min).", "frecuencia": "3 veces por semana"}
        ],
        "ALTO": [
            {"accion": "Meditación en silencio o Mindfulness avanzado.", "frecuencia": "Diario (20 min)"},
            {"accion": "Voluntariado o acto de servicio desinteresado.", "frecuencia": "Mensual"},
            {"accion": "Retiro o desconexión total.", "frecuencia": "Trimestral"}
        ]
    }
}

# --- FUNCIONES DE LÓGICA ---
def leer_informe_anterior_seguro(file_obj):
    try:
        doc = Document(file_obj)
        full_text = " ".join([p.text for p in doc.paragraphs])
        patron = re.search(r"\[DATA\]COG:(\d+);FIS:(\d+);ESP:(\d+)\[END\]", full_text)
        if patron:
            return {
                "PILAR COGNITIVO": int(patron.group(1)),
                "PILAR FÍSICO": int(patron.group(2)),
                "PILAR ESPIRITUAL": int(patron.group(3))
            }
        return None
    except:
        return None

def generar_analisis_avanzado(scores_act, scores_prev):
    """Genera un análisis textual más profundo y detallado."""
    
    # 1. Calcular Porcentajes
    stats = {}
    for pilar, puntos in scores_act.items():
        total = sum(puntos)
        maximo = len(puntos) * 4
        stats[pilar] = (total / maximo) * 100

    promedio_global = sum(stats.values()) / 3
    
    # Identificar Fortaleza y Debilidad
    ordenados = sorted(stats.items(), key=lambda x: x[1])
    debilidad_nombre, debilidad_valor = ordenados[0]
    fortaleza_nombre, fortaleza_valor = ordenados[-1]
    
    # --- REDACCIÓN DEL INFORME ---
    texto = ""
    
    # Párrafo 1: Estado General
    texto += "RESUMEN DEL ESTADO ACTUAL:\n"
    if promedio_global < 50:
        texto += "Tu diagnóstico indica un estado de alerta o 'modo supervivencia'. Los niveles de energía son bajos en áreas críticas, lo que puede manifestarse como agotamiento crónico o desmotivación. Es prioritario detenerse y recargar antes de exigir más productividad.\n"
    elif promedio_global < 75:
        texto += "Te encuentras en una fase funcional y operativa. Tienes recursos para responder a las demandas diarias, pero es probable que sientas que no estás explotando todo tu potencial o que llegas con lo justo al fin de semana.\n"
    else:
        texto += "Presentas un estado de alto rendimiento y bienestar integral. Existe una sólida coherencia entre tu energía física, mental y espiritual, lo que te permite desplegar tu potencial con fluidez.\n"

    # Párrafo 2: Análisis de Equilibrio (Fortalezas y Debilidades)
    texto += "\nANÁLISIS DE EQUILIBRIO:\n"
    if (fortaleza_valor - debilidad_valor) < 15:
        texto += "Tus pilares se encuentran notablemente equilibrados. Esto es una excelente señal de estabilidad, ya que ninguna área está drenando excesivamente a las otras.\n"
    else:
        texto += f"Existe una descompensación importante. Mientras tu {fortaleza_nombre} es tu gran motor ({fortaleza_valor:.0f}%), el {debilidad_nombre} ({debilidad_valor:.0f}%) está actuando como un freno que limita tu avance general. Enfócate en elevar este último para desbloquear tu energía.\n"

    # Párrafo 3: Evolución (Si hay historial)
    if scores_prev:
        texto += "\nTENDENCIA EVOLUTIVA:\n"
        mejoras = 0
        total_diff = 0
        for pilar in scores_act.keys():
            act = sum(scores_act[pilar])
            prev = scores_prev.get(pilar, 0)
            diff = act - prev
            total_diff += diff
            if diff > 0: mejoras += 1
        
        if mejoras == 3:
            texto += "¡Excelente trayectoria! Has logrado mejoras en los tres pilares simultáneamente. Tu estrategia de bienestar está dando resultados sólidos."
        elif total_diff > 0:
            texto += "El balance general es positivo. Aunque algunos indicadores se han mantenido, la tendencia global es de crecimiento respecto a la medición anterior."
        elif total_diff < 0:
            texto += "Se observa un retroceso en los indicadores generales. Es importante revisar qué hábitos o rutinas se han descuidado recientemente para retomar el rumbo."
        else:
            texto += "Tu estado se ha mantenido estable. Estás consolidando tu nivel actual, lo cual es una buena base para plantearse nuevos desafíos de mejora."
            
    return texto

def set_cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def crear_informe_final_v8(scores_act, scores_prev, nombre, img_stream, resumen_txt, num_reporte):
    doc = Document()
    
    # MARCA INVISIBLE
    s_str = ";".join([f"{k[:3]}:{sum(v)}" for k, v in scores_act.items()])
    marca = f"[DATA]{s_str}[END]"
    run = doc.add_paragraph().add_run(marca)
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(1)

    # HEADER
    tit = doc.add_heading('Informe de Bienestar Integral', 0)
    tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha = datetime.now().strftime('%d-%m-%Y')
    p.add_run(f"Fecha: {fecha}  |  Evaluado: {nombre}  |  Reporte N°: {num_reporte}").bold = True

    # 1. RESUMEN EJECUTIVO (AMPLIADO)
    doc.add_heading('1. Análisis Ejecutivo', 1)
    doc.add_paragraph(resumen_txt)

    # 2. GRÁFICO
    doc.add_heading('2. Visualización de Resultados', 1)
    doc.add_picture(img_stream, width=Inches(6.0))
    doc.add_page_break()

    # 3. PLAN DE MEJORA
    doc.add_heading('3. Estrategia de Mejora Personalizada', 1)
    
    pilares = ["PILAR COGNITIVO", "PILAR FÍSICO", "PILAR ESPIRITUAL"]
    for p in pilares:
        pts = sum(scores_act[p])
        max_p = len(scores_act[p])*4
        pct = (pts/max_p)*100
        nivel = "BAJO" if pct < 50 else "MEDIO" if pct < 75 else "ALTO"
        
        doc.add_heading(f"{p.title()} (Nivel {nivel})", 2)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Acción Recomendada"
        hdr[1].text = "Periodicidad Sugerida"
        
        for cell in hdr:
            set_cell_bg(cell, "EAEAEA")
            cell.paragraphs[0].runs[0].font.bold = True

        acciones = db_plan_accion[p][nivel]
        for acc in acciones:
            row = table.add_row().cells
            row[0].text = acc["accion"]
            row[1].text = acc["frecuencia"]
        doc.add_paragraph("")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio

# --- INTERFAZ DE STREAMLIT ---

# BARRA LATERAL (Sidebar) para el Contador y Logo
with st.sidebar:
    st.header("Centro de Control")
    st.info("Bienvenido al sistema de diagnóstico Dínamo.")
    
    # Mostrar contador actual
    total_reps = leer_contador_actual()
    st.metric(label="Informes Generados", value=total_reps)
    st.markdown("---")
    st.caption("Sistema v8.0 - Análisis Avanzado")

# ÁREA PRINCIPAL
st.title("📊 Diagnóstico de Bienestar Dínamo")
st.markdown("""
Esta herramienta analiza tus tres pilares fundamentales (Cognitivo, Físico y Espiritual) 
y genera un **informe completo con inteligencia analítica** y recomendaciones personalizadas.
""")

col1, col2 = st.columns(2)
with col1:
    nombre = st.text_input("Nombre completo")
with col2:
    uploaded_file = st.file_uploader("📂 ¿Tienes un informe anterior? (Opcional)", type="docx")

# Cuestionario
scores_actual = {"PILAR COGNITIVO": [], "PILAR FÍSICO": [], "PILAR ESPIRITUAL": []}

st.divider()

with st.form("encuesta_form"):
    for pilar, preguntas in db_encuesta.items():
        st.subheader(f"🧠 {pilar}" if "COGNITIVO" in pilar else f"💪 {pilar}" if "FÍSICO" in pilar else f"✨ {pilar}")
        for p in preguntas:
            respuesta = st.radio(
                p["pregunta"], 
                p["opciones"], 
                index=None, 
                key=p["pregunta"]
            )
            if respuesta:
                idx = p["opciones"].index(respuesta) + 1
                scores_actual[pilar].append(idx)
            else:
                scores_actual[pilar].append(0)
        st.markdown("<br>", unsafe_allow_html=True)
    
    submitted = st.form_submit_button("Generar Informe Completo", type="primary")

if submitted:
    respuestas_planas = [item for sublist in scores_actual.values() for item in sublist]
    if 0 in respuestas_planas:
        st.error("⚠️ Por favor responde todas las preguntas para obtener un análisis preciso.")
    else:
        # 1. ACTUALIZAR CONTADOR
        nuevo_total = obtener_y_actualizar_contador()
        
        # 2. PROCESAR
        scores_prev = None
        if uploaded_file:
            scores_prev = leer_informe_anterior_seguro(uploaded_file)
            if scores_prev:
                st.toast("Historial cargado correctamente", icon="✅")
            else:
                st.toast("No se detectó historial válido. Se hará un informe base.", icon="ℹ️")

        # 3. VISUALIZACIÓN
        st.divider()
        st.subheader("Resultados Preliminares")
        
        fig, axes = plt.subplots(1, 3, figsize=(15, 5))
        colores = ['#4e79a7', '#f28e2b', '#e15759']
        pilares = ["PILAR COGNITIVO", "PILAR FÍSICO", "PILAR ESPIRITUAL"]
        
        for i, p in enumerate(pilares):
            ax = axes[i]
            act = sum(scores_actual[p])
            max_p = len(scores_actual[p])*4
            
            if scores_prev:
                prev = scores_prev.get(p, 0)
                ax.bar(["Anterior", "Actual"], [prev, act], color=['#bdc3c7', colores[i]])
                diff = act - prev
                if diff != 0:
                    c = 'green' if diff > 0 else 'red'
                    ax.annotate(f"{diff:+}", xy=(1, act), xytext=(0, prev),
                                arrowprops=dict(arrowstyle="->", color=c, lw=2, fontsize=12, fontweight='bold'))
            else:
                ax.bar(["Obtenido", "Máximo"], [act, max_p], color=[colores[i], '#eaeaea'])
            
            pct = (act/max_p)*100
            ax.set_title(f"{p}\n{pct:.0f}%")
            ax.set_ylim(0, max_p*1.2)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
        
        st.pyplot(fig)
        
        # 4. GENERAR DOCUMENTO
        img_buf = io.BytesIO()
        plt.savefig(img_buf, format='png', bbox_inches='tight')
        img_buf.seek(0)
        
        analisis_texto = generar_analisis_avanzado(scores_actual, scores_prev)
        
        docx_file = crear_informe_final_v8(
            scores_actual, 
            scores_prev, 
            nombre if nombre else "Usuario", 
            img_buf, 
            analisis_texto,
            nuevo_total
        )
        
        st.success("¡Informe generado y listo para descargar!")
        st.info(f"Este es el reporte número {nuevo_total} generado por el sistema.")
        
        fecha_str = datetime.now().strftime('%d-%m-%Y')
        nombre_clean = nombre.replace(' ', '_') if nombre else "Usuario"
        
        st.download_button(
            label="⬇️ Descargar Informe PDF/Word",
            data=docx_file,
            file_name=f"Informe_Dinamo_{nombre_clean}_{fecha_str}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )